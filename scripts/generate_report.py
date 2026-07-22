"""Generate use case diagram (PNG) + comprehensive DOCX report for supervisor."""

import os, sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
os.environ.setdefault("DJANGO_SETTINGS_MODULE", "config.settings")
import django; django.setup()

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyArrowPatch
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = Path(__file__).parent
DIAGRAM_PATH    = OUT_DIR / "use_case_diagram.png"
PIPELINE_PATH   = OUT_DIR / "pipeline_comparison.png"
CODEBASE_PATH   = OUT_DIR / "codebase_map.png"
DOCX_PATH = Path(__file__).resolve().parent.parent / "SUPERVISOR_REPORT.docx"


# ---------------------------------------------------------------------------
# 1. Use Case Diagram
# ---------------------------------------------------------------------------

def draw_actor(ax, x, y, label, color="#2c3e50"):
    """Draw a UML stick-figure actor."""
    # Head
    head = plt.Circle((x, y + 0.18), 0.07, color=color, zorder=5)
    ax.add_patch(head)
    # Body
    ax.plot([x, x], [y + 0.11, y - 0.10], color=color, lw=1.8, zorder=5)
    # Arms
    ax.plot([x - 0.12, x + 0.12], [y + 0.02, y + 0.02], color=color, lw=1.8, zorder=5)
    # Legs
    ax.plot([x, x - 0.10], [y - 0.10, y - 0.25], color=color, lw=1.8, zorder=5)
    ax.plot([x, x + 0.10], [y - 0.10, y - 0.25], color=color, lw=1.8, zorder=5)
    # Label
    ax.text(x, y - 0.38, label, ha="center", va="top", fontsize=7.5,
            fontweight="bold", color=color, wrap=True,
            multialignment="center")


def draw_usecase(ax, x, y, label, w=0.52, h=0.16, color="#2980b9", fontsize=7):
    ellipse = mpatches.Ellipse((x, y), w, h, color=color, alpha=0.15, zorder=3)
    ax.add_patch(ellipse)
    ellipse2 = mpatches.Ellipse((x, y), w, h, fill=False,
                                 edgecolor=color, lw=1.2, zorder=4)
    ax.add_patch(ellipse2)
    # Wrap long labels
    words = label.split()
    if len(words) > 3:
        mid = len(words) // 2
        label = " ".join(words[:mid]) + "\n" + " ".join(words[mid:])
    ax.text(x, y, label, ha="center", va="center", fontsize=fontsize,
            color="#1a252f", zorder=5, multialignment="center")


def draw_line(ax, x1, y1, x2, y2, color="#7f8c8d"):
    ax.plot([x1, x2], [y1, y2], color=color, lw=0.9, zorder=2, alpha=0.7)


def make_use_case_diagram():
    fig, ax = plt.subplots(figsize=(16, 11))
    ax.set_xlim(0, 16)
    ax.set_ylim(0, 11)
    ax.axis("off")
    fig.patch.set_facecolor("#f8f9fa")

    # System boundary
    sys_box = mpatches.FancyBboxPatch(
        (2.5, 0.5), 11, 10,
        boxstyle="round,pad=0.1",
        linewidth=2, edgecolor="#2c3e50",
        facecolor="white", zorder=1
    )
    ax.add_patch(sys_box)
    ax.text(8.0, 10.25, "ClinNoteRAG System", ha="center", va="center",
            fontsize=13, fontweight="bold", color="#2c3e50")

    # ── Actors (left side) ──────────────────────────────────────────────────
    actors_left = [
        (1.2, 8.5, "Medical\nStudent"),
        (1.2, 5.5, "Faculty /\nGrader"),
        (1.2, 2.5, "System\nAdmin"),
    ]
    for x, y, lbl in actors_left:
        draw_actor(ax, x, y, lbl, color="#2c3e50")

    # ── Actors (right side — external systems) ──────────────────────────────
    actors_right = [
        (14.8, 8.0, "CAIR\nLLM Server"),
        (14.8, 4.5, "ChromaDB\nVector Store"),
        (14.8, 1.8, "NBME\nDataset"),
    ]
    for x, y, lbl in actors_right:
        draw_actor(ax, x, y, lbl, color="#8e44ad")

    # ── Use Cases ───────────────────────────────────────────────────────────
    # Student use cases
    student_ucs = [
        (5.5, 9.2, "Submit Patient Note"),
        (5.5, 8.2, "View Coverage Score"),
        (5.5, 7.2, "View Evidence Quotes"),
        (5.5, 6.2, "View Per-Concept Feedback"),
    ]
    for x, y, lbl in student_ucs:
        draw_usecase(ax, x, y, lbl, color="#2980b9")

    # Faculty/Grader use cases
    grader_ucs = [
        (8.0, 5.8, "Review Student Submissions"),
        (8.0, 4.8, "View Evaluation Dashboard"),
        (8.0, 3.8, "Export Results / Metrics"),
        (5.5, 4.8, "Manage Cases & Concepts"),
    ]
    for x, y, lbl in grader_ucs:
        draw_usecase(ax, x, y, lbl, color="#27ae60")

    # Admin use cases
    admin_ucs = [
        (5.5, 3.2, "Manage Users & Roles"),
        (5.5, 2.2, "Ingest Concepts to ChromaDB"),
        (5.5, 1.2, "Run Evaluation Harness"),
    ]
    for x, y, lbl in admin_ucs:
        draw_usecase(ax, x, y, lbl, color="#e67e22")

    # External system use cases
    ext_ucs = [
        (11.5, 9.0, "Grade Note (No-RAG)"),
        (11.5, 8.0, "Grade with Concepts (Naive RAG)"),
        (11.5, 7.0, "Evaluate per Concept (Agentic RAG)"),
        (11.5, 4.8, "Store Concept Embeddings"),
        (11.5, 3.8, "Retrieve Concept Synonyms"),
        (11.5, 1.8, "Provide Annotated Notes"),
        (11.5, 1.0, "Provide Rubric Features"),
    ]
    for x, y, lbl in ext_ucs:
        draw_usecase(ax, x, y, lbl, color="#8e44ad")

    # ── Association lines ───────────────────────────────────────────────────
    # Student → student_ucs
    for _, y, _ in student_ucs:
        draw_line(ax, 1.55, 8.5, 5.24, y)

    # Grader → grader_ucs
    for x, y, _ in grader_ucs:
        draw_line(ax, 1.55, 5.5, x - 0.26, y)

    # Admin → admin_ucs
    for _, y, _ in admin_ucs:
        draw_line(ax, 1.55, 2.5, 5.24, y)

    # CAIR → ext LLM ucs
    for _, y, _ in ext_ucs[:3]:
        draw_line(ax, 14.45, 8.0, 11.76, y)

    # ChromaDB → ext chroma ucs
    for _, y, _ in ext_ucs[3:5]:
        draw_line(ax, 14.45, 4.5, 11.76, y)

    # NBME → ext nbme ucs
    for _, y, _ in ext_ucs[5:]:
        draw_line(ax, 14.45, 1.8, 11.76, y)

    # Legend
    legend_items = [
        mpatches.Patch(color="#2980b9", alpha=0.5, label="Student interactions"),
        mpatches.Patch(color="#27ae60", alpha=0.5, label="Faculty / Grader interactions"),
        mpatches.Patch(color="#e67e22", alpha=0.5, label="Admin interactions"),
        mpatches.Patch(color="#8e44ad", alpha=0.5, label="External system interactions"),
    ]
    ax.legend(handles=legend_items, loc="lower right", fontsize=8,
              framealpha=0.9, bbox_to_anchor=(0.99, 0.01))

    ax.set_title("ClinNoteRAG — UML Use Case Diagram", fontsize=15,
                 fontweight="bold", color="#2c3e50", pad=10)

    plt.tight_layout()
    plt.savefig(DIAGRAM_PATH, dpi=180, bbox_inches="tight",
                facecolor=fig.get_facecolor())
    plt.close()
    print(f"Diagram saved → {DIAGRAM_PATH}")


# ---------------------------------------------------------------------------
# 2. DOCX helpers
# ---------------------------------------------------------------------------

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_heading(doc, text, level=1, color="1F3864"):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = RGBColor.from_string(color)
    return p


def add_table(doc, headers, rows, header_color="1F3864"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        set_cell_bg(cell, header_color)
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for r_idx, row in enumerate(rows):
        tr = table.rows[r_idx + 1]
        bg = "F2F2F2" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row):
            cell = tr.cells[c_idx]
            cell.text = str(val)
            set_cell_bg(cell, bg)
            cell.paragraphs[0].runs[0].font.size = Pt(8.5)

    doc.add_paragraph()
    return table


# ---------------------------------------------------------------------------
# 3. Build DOCX
# ---------------------------------------------------------------------------

def build_docx():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.1)
        section.right_margin = Inches(1.1)

    # ── Cover Page ──────────────────────────────────────────────────────────
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Agentic RAG for Clinical Concept Coverage Assessment\nof Medical Student Patient Notes")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor.from_string("1F3864")

    doc.add_paragraph()
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("Supervisor Report — System Design, Pipeline Architecture & Evaluation Results").font.size = Pt(12)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run(
        "Supervisor: Dr. Elif Ak\nAssistant Professor, Department of ECE\n\n"
        "Golam Sarwar Md. Mursalin    ID: 202481996\n"
        "S M Ziauddin    ID: 202482114\n\n"
        "Department of Electrical and Computer Engineering\n"
        "Memorial University of Newfoundland\n"
        "July 2026"
    )
    meta_run.font.size = Pt(11)

    doc.add_page_break()

    # ── Section 1: System Overview ──────────────────────────────────────────
    add_heading(doc, "1. System Overview", level=1)
    doc.add_paragraph(
        "ClinNoteRAG is a zero-shot Agentic RAG system that automatically grades "
        "medical student free-text patient notes against NBME clinical rubric concepts. "
        "Given a patient history note and a set of rubric features for a clinical case, "
        "the system determines for each concept whether it is present or absent in the note, "
        "and extracts the exact evidence phrase supporting that verdict. "
        "The system was evaluated on 40,000+ expert physician annotations across 10 NBME "
        "clinical cases, achieving F1 = 0.942 in the No-RAG configuration — matching or "
        "exceeding supervised fine-tuned transformer models without any labelled training data."
    )

    # ── Section 2: Entities ─────────────────────────────────────────────────
    add_heading(doc, "2. System Entities", level=1)
    doc.add_paragraph(
        "The following entities interact with the ClinNoteRAG system:"
    )

    add_heading(doc, "2.1 Human Actors", level=2)
    add_table(doc,
        headers=["Actor", "Role", "Primary Interactions"],
        rows=[
            ["Medical Student", "End user submitting patient notes for grading",
             "Submit note, view coverage score, view evidence quotes, view per-concept feedback"],
            ["Faculty / Grader", "Instructor reviewing student submissions",
             "Review submissions, view evaluation dashboard, export results and metrics, manage cases"],
            ["System Admin", "Technical operator maintaining the platform",
             "Manage users and roles, ingest NBME concepts to ChromaDB, run evaluation harness, configure system"],
        ]
    )

    add_heading(doc, "2.2 External Systems", level=2)
    add_table(doc,
        headers=["External System", "Purpose", "Interactions"],
        rows=[
            ["CAIR LLM Server\n(granite-4.1-30b)", "Institutional LLM server via LiteLLM (OpenAI-compatible API)",
             "Grades patient notes in No-RAG, Naive RAG, and Agentic RAG modes"],
            ["ChromaDB Vector Store", "Local vector database storing concept embeddings",
             "Stores concept embeddings at ingestion; retrieves concept synonyms during evaluation"],
            ["NBME Dataset\n(Kerbie Addis / Kaggle)", "Source of patient notes and expert physician annotations",
             "Provides 2,839 annotated notes and 143 rubric features across 10 clinical cases"],
        ]
    )

    # ── Section 3: Use Case Diagram ─────────────────────────────────────────
    add_heading(doc, "3. Use Case Diagram", level=1)
    doc.add_paragraph(
        "The diagram below shows all system entities and their interactions. "
        "Human actors appear on the left; external systems appear on the right. "
        "Use cases are grouped by colour: blue (student), green (faculty/grader), "
        "orange (admin), and purple (external systems)."
    )
    doc.add_picture(str(DIAGRAM_PATH), width=Inches(6.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ── Section 4: Pipeline Diagrams ────────────────────────────────────────
    add_heading(doc, "4. Pipeline Architecture", level=1)
    doc.add_paragraph(
        "The diagrams below show exactly how each of the three evaluation strategies "
        "processes a patient note — where data flows, which components are called, "
        "and how they differ from each other."
    )

    add_heading(doc, "4.1 Strategy Comparison", level=2)
    doc.add_paragraph(
        "All three strategies share the same input (patient note + NBME rubric concepts) "
        "and produce the same output (present/absent verdict + evidence quote per concept). "
        "The difference is in how — and whether — they use ChromaDB retrieval and how "
        "many LLM calls they make per note."
    )
    doc.add_picture(str(PIPELINE_PATH), width=Inches(6.4))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    add_table(doc,
        headers=["Strategy", "ChromaDB Calls / Note", "LLM Calls / Note", "What is Retrieved"],
        rows=[
            ["No-RAG", "0", "1", "Nothing — LLM uses pre-trained knowledge only"],
            ["Naive RAG", "1 (batch)", "1", "Synonym lists for all 13–17 concepts at once"],
            ["Agentic RAG", "13–17 (per concept)", "13–17", "One concept's synonyms per tool call"],
        ]
    )

    add_heading(doc, "4.2 How Each Strategy Works", level=2)

    add_heading(doc, "No-RAG", level=3)
    doc.add_paragraph(
        "The patient note and a plain list of concept names are placed directly into a single "
        "prompt. No ChromaDB is involved. The LLM relies entirely on its pre-trained medical "
        "knowledge to recognise whether each concept is mentioned in the note. "
        "One API call produces verdicts for all 13–17 concepts simultaneously. "
        "This is the fastest and cheapest strategy — and in our ablation study, the highest F1 "
        "(0.942), because Granite-4.1-30b already knows medical vocabulary as well as or better "
        "than our synonym lists."
    )

    add_heading(doc, "Naive RAG", level=3)
    doc.add_paragraph(
        "Before calling the LLM, the system queries ChromaDB to retrieve synonym lists for "
        "all concepts belonging to this clinical case in a single batch query. These synonym "
        "lists (e.g. 'shortness of breath | SOB | dyspnea | breathlessness') are inserted "
        "into the prompt alongside the note. The LLM then makes one API call. "
        "The hypothesis was that synonym expansion would help the LLM catch edge cases — "
        "but because the LLM already knows these synonyms, it added noise rather than "
        "signal, resulting in F1=0.923, below No-RAG."
    )

    add_heading(doc, "Agentic RAG", level=3)
    doc.add_paragraph(
        "A Pydantic AI agent receives the note and a list of feature numbers. For each concept, "
        "it autonomously calls the retrieve_concept_info() tool, which queries ChromaDB and "
        "returns that concept's synonyms. The agent then decides present/absent and extracts "
        "evidence — before moving to the next concept. This produces 13–17 LLM calls per note "
        "with a growing conversation history. The multi-step reasoning trace makes each "
        "verdict fully auditable, but compounding errors across 13–17 calls resulted in the "
        "lowest F1 (0.789) of the three completed strategies."
    )

    doc.add_page_break()

    # ── Section 5: Codebase Map ──────────────────────────────────────────────
    add_heading(doc, "5. Codebase Architecture", level=1)
    doc.add_paragraph(
        "The diagram below maps every file in the ClinNoteRAG codebase to its role "
        "in the pipeline — from raw NBME data files through ingestion, vector storage, "
        "RAG services, evaluation harness, and the Django web application."
    )
    doc.add_picture(str(CODEBASE_PATH), width=Inches(6.4))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    add_heading(doc, "Key Files", level=2)
    add_table(doc,
        headers=["File / Directory", "Role"],
        rows=[
            ["scripts/ingest_concepts.py", "Reads NBME features file, embeds 143 concepts, stores in ChromaDB (nbme_concepts)"],
            ["services/embedder.py", "Wrapper around sentence-transformers all-MiniLM-L6-v2 for embedding generation"],
            ["services/no_rag.py", "assess_note_no_rag() — zero retrieval, one LLM call, concept names only"],
            ["services/naive_rag.py", "assess_note_naive() — batch ChromaDB retrieval, one LLM call with synonym lists"],
            ["services/agent.py", "assess_note() — Pydantic AI agent with retrieve_concept_info tool, 13-17 LLM calls"],
            ["scripts/evaluate.py", "Evaluation harness: --strategy flag selects pipeline, saves results.csv + metrics.csv incrementally"],
            ["apps/assessment/views.py", "Django views: student note submission, grader dashboard, admin panel"],
            ["config/settings.py", "Central config: NBME data paths, ChromaDB path, CAIR LLM credentials"],
        ]
    )

    doc.add_page_break()

    # ── Section 6: Why No-RAG Wins & Planned Improvements ───────────────────
    add_heading(doc, "6. Current Finding: Why No-RAG Outperforms RAG", level=1)
    doc.add_paragraph(
        "After running all three strategies, an unexpected result emerged: No-RAG (F1=0.942) "
        "outperforms both Naive RAG (F1=0.923) and Agentic RAG (F1=0.789). "
        "This section explains why, and outlines the improvements planned for the next phase."
    )

    add_heading(doc, "6.1 Root Cause", level=2)
    doc.add_paragraph(
        "The current ChromaDB stores concept synonym lists sourced from the NBME features file, for example:"
    )
    p = doc.add_paragraph(style="No Spacing")
    p.add_run("    shortness of breath  |  SOB  |  dyspnea  |  breathlessness  |  difficulty breathing").font.name = "Courier New"
    doc.add_paragraph()
    doc.add_paragraph(
        "Granite-4.1-30b is a 30-billion parameter model trained on vast medical literature. "
        "It already knows every one of those synonyms. Retrieving them from ChromaDB gives the "
        "LLM no new information — it adds noise to the prompt without adding knowledge. "
        "RAG only outperforms a zero-shot LLM when the retrieved content contains information "
        "the model genuinely cannot have learned during training."
    )

    add_heading(doc, "6.2 Planned Improvements", level=2)
    doc.add_paragraph(
        "Three targeted improvements are planned. Each addresses a specific weakness identified "
        "from the ablation results."
    )

    # Option 1
    add_heading(doc, "Option 1 — Retrieve Expert-Labeled Similar Notes (Highest Impact)", level=3)
    doc.add_paragraph(
        "Instead of storing concept synonym lists in ChromaDB, embed the 2,839 expert-annotated "
        "patient notes and store them with their physician labels. For any new student note, "
        "retrieve the 3–5 most similar annotated notes and show the LLM how a physician labeled "
        "each concept in those similar cases."
    )
    p = doc.add_paragraph(style="No Spacing")
    p.add_run(
        "    New note: \"patient says heart racing, no chest pain, mother had cardiac issues\"\n\n"
        "    Retrieved similar note (cosine similarity 0.91):\n"
        "      \"17yo male, palpitations x3 months, mother with MI at 52\"\n"
        "      → family history of MI: PRESENT — \"mother with MI\"\n"
        "      → palpitations: PRESENT — \"palpitations x3 months\"\n"
        "      → chest pain: ABSENT\n\n"
        "    Now evaluate the new note using these as examples."
    ).font.name = "Courier New"
    doc.add_paragraph()
    doc.add_paragraph(
        "This is genuine RAG — the LLM receives expert grading decisions on cases it has never "
        "seen. To prevent data leakage, an 80/20 train/test split is used: 2,271 notes form "
        "the retrieval index; 568 held-out notes are used for evaluation."
    )

    add_heading(doc, "Option 2 — Fix Agentic RAG Architecture (On-Demand Retrieval)", level=3)
    doc.add_paragraph(
        "The current Agentic RAG makes 13–17 sequential API calls per note — one forced tool "
        "call per concept — with conversation history growing after each call. This compounds "
        "errors and inflates cost. The improved architecture makes one API call per note but "
        "allows the agent to call retrieve_similar_notes(concept) only when it is uncertain, "
        "making retrieval targeted rather than mechanical."
    )

    add_heading(doc, "6.3 Expected Impact", level=2)
    add_table(doc,
        headers=["Configuration", "What Changes", "Expected Effect"],
        rows=[
            ["Current No-RAG", "Baseline — LLM uses pre-trained knowledge only", "F1 = 0.942"],
            ["Current Naive RAG", "ChromaDB synonym lists (LLM already knows these)", "F1 = 0.923 (worse)"],
            ["Option 1 — Similar Note RAG", "Expert-labeled notes retrieved per query", "F1 >> 0.942 (target)"],
            ["Option 1+2 Combined", "Full improved pipeline", "Best possible F1"],
        ]
    )

    add_heading(doc, "6.4 Research Narrative", level=2)
    doc.add_paragraph(
        "The finding that No-RAG outperforms RAG is itself a contribution: it demonstrates that "
        "synonym-based retrieval fails when the LLM already possesses that knowledge. The "
        "subsequent improvement — retrieving expert-annotated similar cases — will show that RAG "
        "only adds value when it provides information the model genuinely lacks. This nuanced "
        "finding is more publishable than simply showing RAG helps, because it identifies "
        "precisely when and why retrieval matters in medical education assessment."
    )

    add_heading(doc, "6.5 Implementation Roadmap", level=2)
    add_table(doc,
        headers=["Step", "Task", "Rationale"],
        rows=[
            ["1", "Implement Option 1 — embed notes, 80/20 split, similar-note retrieval",
             "Highest impact; directly answers supervisor's concern about LLM pre-training dominating"],
            ["2", "Implement Option 2 — on-demand agentic retrieval",
             "Architecture quality and cost reduction; combine with Option 1"],
        ]
    )

    doc.add_page_break()

    # ── Section 5: Competitive Analysis ─────────────────────────────────────
    add_heading(doc, "7. Competitive Analysis", level=1)
    doc.add_paragraph(
        "This section compares ClinNoteRAG against existing systems in automated "
        "medical note grading, clinical NLP tools, commercial platforms, and LLM-based approaches."
    )

    add_heading(doc, "7.1 Task Clarification", level=2)
    doc.add_paragraph(
        "The Kaggle NBME 2022 competition framed this as a span extraction task "
        "(find character offsets of evidence in the note). ClinNoteRAG frames it as "
        "binary classification with evidence extraction (present/absent + quoted text). "
        "These are related but not identical — Kaggle scores are not directly comparable "
        "to our F1, but they are the closest academic precedent."
    )

    add_heading(doc, "7.2 Kaggle NBME 2022 Competition (Closest Academic Precedent)", level=2)
    add_table(doc,
        headers=["System", "Method", "F1", "Key Limitation"],
        rows=[
            ["1st place ensemble", "DeBERTa-v3-large + v2-xxlarge, task-adaptive pretraining", "~0.890",
             "Fully supervised — requires thousands of labelled span annotations; no reasoning"],
            ["General top solutions", "DeBERTa token classification / span extraction", "0.85–0.89",
             "Brittle — requires retraining for new cases; outputs character offsets only"],
            ["INCITE (published academic)", "Fine-tuned transformer, span-based", "0.890",
             "Not zero-shot; per-case supervised training required"],
            ["Two-phase LLM (JMIR 2025)", "Section detection → LLM feature extraction", "0.968–0.983",
             "Tested on 700 notes only — not validated at scale; few-shot, not zero-shot"],
        ]
    )

    add_heading(doc, "7.3 Traditional Clinical NLP Tools", level=2)
    add_table(doc,
        headers=["System", "Method", "F1", "Key Limitation"],
        rows=[
            ["MetaMap (NIH)", "Rule-based + UMLS Metathesaurus lookup", "~0.88",
             "Dictionary lookup only; cannot handle paraphrases, negation, or student misspellings"],
            ["cTAKES (Apache)", "Rule-based + ML + SNOMED-CT", "~0.89",
             "Outputs general UMLS concepts — not rubric-specific; heavy pipeline engineering needed"],
            ["CLAMP", "ML + rules", "0.70 (i2b2) / 0.39 (MIMIC-III)",
             "Significant accuracy drop on realistic messy clinical text"],
            ["MedSpaCy", "spaCy-based, manual rule extensions", "Not benchmarked on NBME",
             "Requires hand-crafting rules per concept set; not generalisable"],
        ]
    )

    add_heading(doc, "7.4 Commercial Medical Education Platforms", level=2)
    add_table(doc,
        headers=["Platform", "Free-text Note Grading?", "Limitation vs ClinNoteRAG"],
        rows=[
            ["NBME's own scoring", "No — human physician raters only",
             "Expensive, slow (~$50–100/note), cannot scale to thousands of students"],
            ["ExamSoft", "No — MCQ/checklist only",
             "Cannot process narrative notes; no NLP component"],
            ["Aquifer", "No — virtual cases with MCQ",
             "Out of scope for free-text patient note evaluation"],
            ["Osmosis / Amboss", "No — flashcard and question banks",
             "No open-ended assessment capability"],
        ]
    )

    add_heading(doc, "7.5 LLM-Based Assessment Systems", level=2)
    add_table(doc,
        headers=["System", "Task", "Performance", "Limitation"],
        rows=[
            ["GPT-4 (OpenAI)", "De-identification, criteria extraction", "~0.897 F1",
             "Not evaluated on NBME rubric grading; expensive proprietary API"],
            ["MedPaLM 2 (Google)", "USMLE MCQ answering", "86.5% accuracy",
             "MCQ only — no published results on patient note grading"],
            ["ChatGPT for OSCE (2024)", "OSCE verbal checklist scoring", "~80–85% agreement",
             "Structured checklist only; no free-text notes; no evidence extraction"],
            ["i-MedRAG", "Iterative RAG for medical QA", "Improves over RAG baseline",
             "Different task — not patient note grading; not benchmarked on NBME"],
        ]
    )

    add_heading(doc, "7.6 Master Comparison Table", level=2)
    add_table(doc,
        headers=["System", "Zero-shot?", "Evidence Extraction?", "F1", "Labels Evaluated"],
        rows=[
            ["ClinNoteRAG No-RAG (ours)", "Yes", "Yes — quoted text", "0.942", "40,515"],
            ["ClinNoteRAG Naive RAG (ours)", "Yes", "Yes — quoted text", "0.923", "40,453"],
            ["ClinNoteRAG Agentic RAG (ours)", "Yes", "Yes — quoted text", "0.789", "39,577"],
            ["Kaggle 1st place (DeBERTa)", "No — supervised", "Span offsets only", "~0.890", "~14,000"],
            ["INCITE (academic)", "No — supervised", "Span offsets only", "0.890", "~14,000"],
            ["Two-phase LLM (JMIR 2025)", "Few-shot", "Partial", "0.983", "700 only"],
            ["MetaMap / cTAKES", "No — dictionary", "No", "~0.88", "EHR corpora"],
            ["ExamSoft / NBME human scoring", "N/A", "N/A", "N/A (manual)", "Manual"],
            ["GPT-4 (various studies)", "Yes", "Partial", "~0.90", "Small studies"],
        ]
    )

    add_heading(doc, "7.7 Where ClinNoteRAG Stands Out", level=2)
    differentiators = [
        ("Zero-shot at scale",
         "40,000+ labels evaluated with no labelled training data, matching or exceeding "
         "supervised fine-tuned DeBERTa ensembles that required full annotation."),
        ("Interpretable evidence extraction",
         "Returns a quoted text span per concept. No prior system does this on this dataset — "
         "DeBERTa outputs character offsets, not human-readable evidence."),
        ("First systematic ablation on this task",
         "No-RAG vs Naive RAG vs Agentic RAG comparison with 40,000+ labels. Key finding: "
         "LLM pre-training dominates; retrieval adds no measurable F1 benefit on this dataset."),
        ("Free institutional deployment",
         "Runs on MUN CAIR server (granite-4.1-30b via LiteLLM) — no per-note API cost "
         "for production deployment."),
        ("Open market gap",
         "No commercial platform today automatically grades free-text NBME patient notes "
         "against a rubric. ClinNoteRAG directly addresses this gap."),
    ]
    for title, desc in differentiators:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(title + ": ")
        run.bold = True
        run.font.size = Pt(10)
        p.add_run(desc).font.size = Pt(10)

    doc.add_page_break()

    # ── Section 6: Ablation Results Summary ─────────────────────────────────
    add_heading(doc, "8. Ablation Study Results", level=1)
    doc.add_paragraph(
        "Three evaluation strategies were compared across 2,800+ NBME patient notes "
        "and 40,000+ expert physician labels."
    )
    add_table(doc,
        headers=["Strategy", "Precision", "Recall", "F1", "Labels", "LLM Calls / Note"],
        rows=[
            ["No-RAG", "0.9418", "0.9428", "0.9423", "40,515", "1"],
            ["Naive RAG", "0.9065", "0.9406", "0.9233", "40,453", "1"],
            ["Agentic RAG", "0.7110", "0.8852", "0.7886", "39,577", "13–17"],
        ]
    )
    doc.add_paragraph(
        "Key finding: The LLM's pre-trained medical knowledge (No-RAG) is the dominant "
        "factor. Adding synonym retrieval from ChromaDB (Naive RAG) provides no measurable "
        "F1 improvement. The Agentic RAG approach trades accuracy for interpretability — "
        "it produces a full reasoning trace and per-concept evidence quotes, which are "
        "essential for student feedback even at a lower aggregate F1."
    )

    doc.save(DOCX_PATH)
    print(f"DOCX saved → {DOCX_PATH}")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    print("Generating use case diagram...")
    make_use_case_diagram()
    print("Building DOCX report...")
    build_docx()
    print("Done.")
