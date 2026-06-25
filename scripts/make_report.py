"""Generate PROGRESS_REPORT.docx"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

OUT = Path(__file__).parent.parent / "PROGRESS_REPORT.docx"

doc = Document()

for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)


def centered(text, size=12, bold=False, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    return p


def heading(text, level=1):
    return doc.add_heading(text, level=level)


def para(text):
    p = doc.add_paragraph()
    p.add_run(text)
    p.paragraph_format.space_after = Pt(8)
    return p


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            table.rows[r + 1].cells[c].text = val
    doc.add_paragraph()


def screenshot_placeholder(caption):
    """Adds a bordered single-cell table as a screenshot placeholder."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    cell.width = Inches(5.5)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(40)
    p.paragraph_format.space_after = Pt(40)
    run = p.add_run(f"[ Screenshot placeholder: {caption} ]")
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.font.size = Pt(10)
    run.italic = True
    doc.add_paragraph()


# ── Cover Page ─────────────────────────────────────────────────────────────
doc.add_paragraph().paragraph_format.space_after = Pt(48)

centered("Agentic RAG for Clinical Concept Coverage Assessment", size=16, bold=True, space_after=4)
centered("of Medical Student Patient Notes", size=16, bold=True, space_after=32)

centered("Progress Report", size=13, bold=False, space_after=32)

centered("Supervisor: Dr. Elif Ak", size=11, space_after=2)
centered("Assistant Professor, Department of Electrical and Computer Engineering", size=11, space_after=32)

centered("Golam Sarwar Md. Mursalin    ID: 202481996", size=11, space_after=2)
centered("S M Ziauddin    ID: 202482114", size=11, space_after=32)

centered("Department of Electrical and Computer Engineering", size=11, space_after=2)
centered("Memorial University of Newfoundland", size=11, space_after=2)
centered("June 2026", size=11, space_after=0)

doc.add_page_break()

# ── Section 1 ──────────────────────────────────────────────────────────────
heading("What We're Building")
para(
    "ClinNoteRAG is an agentic RAG system that automatically grades medical student patient "
    "notes against NBME rubric concepts. The idea is straightforward — given a free-text note "
    "written by a student and a case ID, the system goes through each clinical concept in the "
    "rubric and decides whether the student mentioned it, along with a quote from the note as "
    "evidence."
)
para(
    "We are evaluating against the NBME USMLE Step 2 dataset from Kaggle, which has 14,300 "
    "expert physician labels across 143 rubric concepts and 10 clinical cases."
)
para(
    "What makes it agentic is that the system does not just dump everything into one prompt. "
    "It uses a Pydantic AI agent with a retrieval tool — for each concept, it fetches the "
    "concept's synonym variants from ChromaDB, then asks Claude to check if any of those "
    "variants appear in the note. This targeted per-concept reasoning is the core research "
    "contribution."
)

# ── Section 2 ──────────────────────────────────────────────────────────────
heading("What's Done So Far (Phase 1)")
para(
    "We have a fully working system end-to-end. The Django project is up and running with "
    "a web dashboard, all 143 NBME concepts are ingested and embedded in ChromaDB, and the "
    "agentic pipeline is complete. We also wrote the evaluation harness that runs the agent "
    "over notes and computes Precision, Recall, and F1 against the expert labels."
)

add_table(
    ["Component", "Status"],
    [
        ["Django project + web dashboard", "Done"],
        ["ChromaDB — 143 concepts ingested", "Done"],
        ["Agentic RAG pipeline (Claude + tool use)", "Done"],
        ["Evaluation harness (P/R/F1)", "Done"],
        ["Smoke test on one note", "Passed"],
        ["Full quantitative evaluation", "Pending"],
    ]
)

para(
    "Stack: Django 5, ChromaDB, Claude Haiku (Anthropic API), Pydantic AI, "
    "sentence-transformers, scikit-learn."
)

# ── Section 3 ──────────────────────────────────────────────────────────────
heading("Web Dashboard")
para(
    "We built a web dashboard using Django to browse the knowledge base and track evaluation "
    "runs. There are three main pages — the main dashboard showing all 10 NBME cases and past "
    "evaluation runs, a concept browser for each case showing how concepts and their synonyms "
    "are stored in ChromaDB, and a run detail page showing per-case F1 breakdown once an "
    "evaluation completes."
)

screenshot_placeholder("Main dashboard — http://127.0.0.1:8000/")
screenshot_placeholder("Concept browser — Case 0 (http://127.0.0.1:8000/case/0/)")

# ── Section 4 ──────────────────────────────────────────────────────────────
heading("Smoke Test Results")
para(
    "We ran the agent on one note — Case 0 (palpitations, 17-year-old male, 13 concepts) — "
    "to verify everything works before the full evaluation. The results looked correct. "
    "A few highlights:"
)

add_table(
    ["Concept", "Verdict", "Evidence from note"],
    [
        ["Family history of MI", "PRESENT", '"father had MI recently"'],
        ["Adderall use", "PRESENT", '"meds: aderol (from a friend)"'],
        ["Shortness of breath", "PRESENT", '"dispnea on exersion and rest"'],
        ["Intermittent symptoms", "PRESENT", '"intermittent for 2 days (lasting 3-4 min)"'],
        ["Chest pressure", "absent", "—"],
    ]
)

para(
    "One thing worth noting — the agent correctly matched \"dispnea\" (a misspelling) to the "
    "Shortness of breath concept because the vector similarity search handled it. This is "
    "exactly the kind of thing a simple keyword match would miss."
)

# ── Section 5 ──────────────────────────────────────────────────────────────
heading("Full Evaluation — Status")
para(
    "We started a 1,000-note evaluation run but had to stop it midway. The agentic approach "
    "makes 13–15 API calls per note (one per concept), which turned out more expensive than "
    "initially estimated — around $9.94 was consumed before we stopped."
)

add_table(
    ["Run", "Notes Target", "Status", "Cost Incurred"],
    [
        ["Agentic RAG (full)", "1,000", "Stopped mid-run", "~$9.94"],
        ["Agentic RAG (limited)", "100", "Planned", "~$4–5"],
    ]
)

para(
    "The fix is to cap the next run at 100 notes, which is enough to report a valid F1 score "
    "and fits the remaining budget. We also improved the evaluation script to save results "
    "incrementally after every note, so a stopped run no longer loses all data."
)

# ── Section 6 ──────────────────────────────────────────────────────────────
heading("Phase 2 Plan")

heading("Ablation Study", level=2)
para(
    "The core research contribution is a three-way ablation study. The goal is to show not "
    "just that the system works, but why the agentic design is better than simpler alternatives. "
    "We will run two additional baselines and compare all three side by side."
)
para(
    "The No-RAG baseline gives the LLM only the concept names with no retrieval at all — no "
    "synonyms, no variant expansion. This shows what the model can do purely from its training "
    "knowledge. The Naive RAG baseline retrieves all concepts for a case at once using a single "
    "semantic search and stuffs them all into one big prompt. It is cheaper and simpler, but "
    "the context gets noisy when many concepts are present. The Agentic RAG system — what we "
    "already built — retrieves each concept individually and reasons about it in isolation, "
    "which should give more precise and grounded verdicts."
)

add_table(
    ["Strategy", "How it works", "Expected result"],
    [
        ["No-RAG", "LLM sees concept names only, no retrieval", "Weakest — misses paraphrases and misspellings"],
        ["Naive RAG", "All concepts retrieved at once, one big prompt", "Middle — retrieval helps but context is noisy"],
        ["Agentic RAG", "Per-concept targeted retrieval + reasoning", "Best — precise, grounded verdicts"],
    ]
)

heading("RAGAS Evaluation", level=2)
para(
    "F1 score tells us whether the final verdict was right or wrong, but it does not say "
    "anything about the quality of the retrieval step itself. A system could get lucky and "
    "produce correct verdicts even if it retrieved the wrong concepts. To catch this, we "
    "will integrate RAGAS — an automated RAG evaluation framework — to measure the retrieval "
    "and generation quality independently."
)
para(
    "Specifically, RAGAS measures four things: context precision (out of everything we "
    "retrieved, how much of it was actually relevant?), context recall (did we retrieve all "
    "the information needed to make the right call?), faithfulness (is the evidence quote "
    "actually supported by the note, or did the LLM hallucinate it?), and answer relevancy "
    "(is the verdict on-topic given the concept being evaluated?). Together these metrics "
    "give us a much more detailed picture of where the system is strong and where it "
    "breaks down — which is exactly what makes a research contribution credible."
)

heading("LLM-as-Judge", level=2)
para(
    "The NBME dataset only has expert annotations for 1,000 out of 42,146 notes. "
    "That leaves over 40,000 notes with no ground truth labels. LLM-as-Judge is our "
    "answer to this: we will run a second, independent Claude call that reads each note "
    "and verdict and decides whether it agrees or disagrees — without seeing the first "
    "agent's reasoning."
)
para(
    "We then measure how closely this second Claude agrees with the expert physician labels "
    "using Cohen's kappa, which is the standard metric for inter-rater agreement in medical "
    "research. If the kappa is high, it means Claude can reliably substitute for a human "
    "grader on unannotated notes — which would make ClinNoteRAG useful far beyond the "
    "labelled subset and gives the system real-world clinical education value."
)
para(
    "This is also a direct comparison to the work by Zheng et al. who showed that GPT-4 "
    "can match human evaluator agreement on open-ended tasks. We are testing whether the "
    "same holds in a structured medical grading context using Claude."
)

# ── Section 7 ──────────────────────────────────────────────────────────────
heading("Next Steps")
for step in [
    "Re-run evaluation with --limit 100 to get the agentic RAG F1 score",
    "Build naive RAG and no-RAG baselines",
    "Run all three strategies and complete the ablation table",
    "Integrate RAGAS metrics for retrieval quality evaluation",
    "Implement LLM-as-Judge and measure Cohen's kappa vs expert labels",
    "Add screenshots to this report once the dashboard evaluation run is complete",
]:
    doc.add_paragraph(step, style="List Number")

doc.save(OUT)
print(f"Saved: {OUT}")
